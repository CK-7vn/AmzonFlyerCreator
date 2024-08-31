import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from PIL import Image, ImageTk
from io import BytesIO
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

class FlyerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Flyer Creator")
        self.shoes_list = []
        self.current_shoe = None
        self.current_shoe_index = None

        # Setup GUI
        self.setup_gui()

    def setup_gui(self):
        # Input section
        tk.Label(self.root, text="Amazon URL:").grid(row=0, column=0, padx=10, pady=10)
        self.url_entry = tk.Entry(self.root, width=50)
        self.url_entry.grid(row=0, column=1, padx=10, pady=10)

        tk.Button(self.root, text="Fetch Details", command=self.fetch_details).grid(row=0, column=2, padx=10, pady=10)
        tk.Button(self.root, text="Clear URL", command=self.clear_url).grid(row=0, column=3, padx=10, pady=10)

        tk.Label(self.root, text="Name:").grid(row=1, column=0, padx=10, pady=10)
        self.name_entry = tk.Entry(self.root, width=50)
        self.name_entry.grid(row=1, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Price:").grid(row=2, column=0, padx=10, pady=10)
        self.price_entry = tk.Entry(self.root, width=50)
        self.price_entry.grid(row=2, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Image URL:").grid(row=3, column=0, padx=10, pady=10)
        self.image_entry = tk.Entry(self.root, width=50)
        self.image_entry.grid(row=3, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Amazon URL:").grid(row=4, column=0, padx=10, pady=10)
        self.amz_url_entry = tk.Entry(self.root, width=50)
        self.amz_url_entry.grid(row=4, column=1, padx=10, pady=10)

        tk.Button(self.root, text="Add to Flyer", command=self.add_to_flyer).grid(row=5, column=0, padx=10, pady=10, columnspan=2)
        tk.Button(self.root, text="Update Selected", command=self.update_manually).grid(row=5, column=2, padx=10, pady=10)

        tk.Label(self.root, text="Edit Selected Entry:").grid(row=6, column=0, padx=10, pady=10, columnspan=4)

        tk.Label(self.root, text="Name:").grid(row=7, column=0, padx=10, pady=10)
        self.edit_name_entry = tk.Entry(self.root, width=50)
        self.edit_name_entry.grid(row=7, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Price:").grid(row=8, column=0, padx=10, pady=10)
        self.edit_price_entry = tk.Entry(self.root, width=50)
        self.edit_price_entry.grid(row=8, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Image URL:").grid(row=9, column=0, padx=10, pady=10)
        self.edit_image_entry = tk.Entry(self.root, width=50)
        self.edit_image_entry.grid(row=9, column=1, padx=10, pady=10)

        tk.Label(self.root, text="Amazon URL:").grid(row=10, column=0, padx=10, pady=10)
        self.edit_amz_entry = tk.Entry(self.root, width=50)
        self.edit_amz_entry.grid(row=10, column=1, padx=10, pady=10)

        self.save_flyer_button = tk.Button(self.root, text="Save Flyer", command=self.save_flyer)
        self.save_flyer_button.grid(row=11, column=0, padx=10, pady=10)

        self.save_flyer_list_button = tk.Button(self.root, text="Save Flyer List", command=self.save_flyer_list)
        self.save_flyer_list_button.grid(row=11, column=1, padx=10, pady=10)

        self.show_flyer_button = tk.Button(self.root, text="Show Flyer Preview", command=self.show_flyer_preview)
        self.show_flyer_button.grid(row=11, column=2, padx=10, pady=10)

        self.delete_button = tk.Button(self.root, text="Delete Selected", command=self.delete_entry)
        self.delete_button.grid(row=11, column=3, padx=10, pady=10)

        self.listbox = tk.Listbox(self.root, width=80, height=15)
        self.listbox.grid(row=12, column=0, columnspan=4, padx=10, pady=10)
        self.listbox.bind('<<ListboxSelect>>', self.on_select)

        # Image label initialization
        self.image_label = tk.Label(self.root)
        self.image_label.grid(row=3, column=2, rowspan=2, padx=10, pady=10)

    def fetch_details(self):
        amazon_url = self.url_entry.get()
        if not amazon_url:
            messagebox.showwarning("Input Error", "Please enter an Amazon URL")
            return

        try:
            response = requests.get(amazon_url)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')

            name = soup.find('span', {'id': 'productTitle'})
            name = name.get_text(strip=True) if name else 'Name not available'

            price_whole = soup.find('span', {'class': 'a-price-whole'})
            price_decimal = soup.find('span', {'class': 'a-price-decimal'})
            price_fraction = soup.find('span', {'class': 'a-price-fraction'})

            price = ''
            if price_whole:
                price += price_whole.get_text(strip=True)
            if price_decimal:
                price += '.' + price_decimal.get_text(strip=True)
            if price_fraction:
                price += price_fraction.get_text(strip=True)
            price = price if price else 'Price not available'

            image_tag = soup.find('img', {'data-a-image-name': 'landingImage'})
            image_url = image_tag['src'] if image_tag else 'Image not available'

            self.current_shoe = {'name': name, 'price': price, 'image_url': image_url, 'link': amazon_url}

            self.name_entry.delete(0, tk.END)
            self.name_entry.insert(0, name)

            self.price_entry.delete(0, tk.END)
            self.price_entry.insert(0, price)

            self.image_entry.delete(0, tk.END)
            self.image_entry.insert(0, image_url)

            self.amz_url_entry.delete(0, tk.END)
            self.amz_url_entry.insert(0, amazon_url)

            if image_url != 'Image not available':
                try:
                    image_response = requests.get(image_url)
                    image_response.raise_for_status()
                    image_data = BytesIO(image_response.content)
                    image = Image.open(image_data)
                    image.thumbnail((100, 100))
                    image_tk = ImageTk.PhotoImage(image)
                    self.image_label.config(image=image_tk, text='')
                    self.image_label.image = image_tk
                except Exception as e:
                    self.image_label.config(text='Image not available.')
                    print(f'Error downloading image: {e}')
            else:
                self.image_label.config(text='Image not available.')

        except Exception as e:
            messagebox.showerror("Fetch Error", f"Failed to fetch details: {e}")

    def clear_url(self):
        self.url_entry.delete(0, tk.END)
        self.name_entry.delete(0, tk.END)
        self.price_entry.delete(0, tk.END)
        self.image_entry.delete(0, tk.END)
        self.amz_url_entry.delete(0, tk.END)
        self.image_label.config(image='', text='')

    def add_to_flyer(self):
        name = self.name_entry.get()
        price = self.price_entry.get()
        image_url = self.image_entry.get()
        link = self.amz_url_entry.get()

        if not (name and price and image_url and link):
            messagebox.showwarning("Input Error", "All fields must be filled")
            return

        self.shoes_list.append({'name': name, 'price': price, 'image_url': image_url, 'link': link})
        self.listbox.insert(tk.END, f"{name} - ${price}")
        self.clear_url()

    def on_select(self, event):
        selected_index = self.listbox.curselection()
        if selected_index:
            index = selected_index[0]
            self.current_shoe_index = index
            self.current_shoe = self.shoes_list[index]
            self.edit_name_entry.delete(0, tk.END)
            self.edit_name_entry.insert(0, self.current_shoe['name'])
            self.edit_price_entry.delete(0, tk.END)
            self.edit_price_entry.insert(0, self.current_shoe['price'])
            self.edit_image_entry.delete(0, tk.END)
            self.edit_image_entry.insert(0, self.current_shoe['image_url'])
            self.edit_amz_entry.delete(0, tk.END)
            self.edit_amz_entry.insert(0, self.current_shoe['link'])

            try:
                image_response = requests.get(self.current_shoe['image_url'])
                image_response.raise_for_status()
                image_data = BytesIO(image_response.content)
                image = Image.open(image_data)
                image.thumbnail((100, 100))
                image_tk = ImageTk.PhotoImage(image)
                self.image_label.config(image=image_tk, text='')
                self.image_label.image = image_tk
            except Exception as e:
                self.image_label.config(text='Image not available.')
                print(f'Error downloading image: {e}')

    def update_manually(self):
        if self.current_shoe_index is None:
            messagebox.showwarning("Select Entry", "Please select an entry to update")
            return

        self.shoes_list[self.current_shoe_index] = {
            'name': self.edit_name_entry.get(),
            'price': self.edit_price_entry.get(),
            'image_url': self.edit_image_entry.get(),
            'link': self.edit_amz_entry.get()
        }

        self.listbox.delete(self.current_shoe_index)
        self.listbox.insert(self.current_shoe_index, f"{self.edit_name_entry.get()} - ${self.edit_price_entry.get()}")
        self.clear_edit_entries()

    def clear_edit_entries(self):
        self.edit_name_entry.delete(0, tk.END)
        self.edit_price_entry.delete(0, tk.END)
        self.edit_image_entry.delete(0, tk.END)
        self.edit_amz_entry.delete(0, tk.END)
        self.image_label.config(image='', text='')

    def delete_entry(self):
        selected_index = self.listbox.curselection()
        if selected_index:
            index = selected_index[0]
            del self.shoes_list[index]
            self.listbox.delete(index)
            self.clear_edit_entries()
            self.current_shoe_index = None
            self.current_shoe = None

    def save_flyer(self):
        if not self.shoes_list:
            messagebox.showwarning("No Data", "No data to save")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                               filetypes=[("Word Documents", "*.docx")],
                                               title="Save Flyer As")
        if not file_path:
            return

        doc = Document()
        doc.add_heading('Flyer', level=1)

        items_per_row = 2
        row_count = len(self.shoes_list) // items_per_row + (1 if len(self.shoes_list) % items_per_row else 0)

        for i in range(row_count):
            table = doc.add_table(rows=1, cols=items_per_row)
            table.style = 'Table Grid'

            for j in range(items_per_row):
                index = i * items_per_row + j
                if index < len(self.shoes_list):
                    shoe = self.shoes_list[index]
                    cell = table.cell(0, j)

                    # Add image
                    try:
                        image_response = requests.get(shoe['image_url'])
                        image_response.raise_for_status()
                        image_data = BytesIO(image_response.content)
                        img = Image.open(image_data)
                        img.thumbnail((150, 150))
                        img.save('temp_image.jpg')
                        cell.paragraphs[0].add_run().add_picture('temp_image.jpg', width=Inches(2))
                    except Exception as e:
                        cell.paragraphs[0].add_run("Image not available.")
                        print(f'Error downloading image: {e}')

                    # Add name and price
                    cell.add_paragraph(shoe['name'])
                    cell.add_paragraph(shoe['price'])

        doc.save(file_path)
        messagebox.showinfo("Saved", "Flyer saved successfully")

    def save_flyer_list(self):
        if not self.shoes_list:
            messagebox.showwarning("No Data", "No data to save")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                               filetypes=[("Word Documents", "*.docx")],
                                               title="Save Flyer List As")
        if not file_path:
            return

        doc = Document()
        doc.add_heading('Flyer List', level=1)

        for shoe in self.shoes_list:
            doc.add_paragraph(f"Name: {shoe['name']}")
            doc.add_paragraph(f"Price: {shoe['price']}")
            doc.add_paragraph(f"Link: {shoe['link']}")
            doc.add_paragraph()

        doc.save(file_path)
        messagebox.showinfo("Saved", "Flyer list saved successfully")

    def show_flyer_preview(self):
        if not self.shoes_list:
            messagebox.showwarning("No Data", "No data to preview")
            return

        preview_window = tk.Toplevel(self.root)
        preview_window.title("Flyer Preview")

        preview_canvas = tk.Canvas(preview_window, width=600, height=400)
        preview_canvas.pack(fill=tk.BOTH, expand=True)

        items_per_row = 2
        row_count = len(self.shoes_list) // items_per_row + (1 if len(self.shoes_list) % items_per_row else 0)
        item_width = 300
        item_height = 300

        for i in range(row_count):
            for j in range(items_per_row):
                index = i * items_per_row + j
                if index < len(self.shoes_list):
                    shoe = self.shoes_list[index]
                    x = j * item_width
                    y = i * item_height

                    # Add image
                    try:
                        image_response = requests.get(shoe['image_url'])
                        image_response.raise_for_status()
                        image_data = BytesIO(image_response.content)
                        img = Image.open(image_data)
                        img.thumbnail((150, 150))
                        img_tk = ImageTk.PhotoImage(img)
                        preview_canvas.create_image(x + 150 // 2, y + 50, image=img_tk)
                        preview_canvas.create_text(x + 150 // 2, y + 170, text=shoe['name'], anchor=tk.N)
                        preview_canvas.create_text(x + 150 // 2, y + 200, text=shoe['price'], anchor=tk.N)
                    except Exception as e:
                        preview_canvas.create_text(x + 150 // 2, y + 50, text='Image not available.')
                        print(f'Error downloading image: {e}')

if __name__ == "__main__":
    root = tk.Tk()
    app = FlyerApp(root)
    root.mainloop()
